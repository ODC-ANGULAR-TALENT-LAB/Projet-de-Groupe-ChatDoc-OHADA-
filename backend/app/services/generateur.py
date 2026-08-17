"""Generateur de documents types, adosse aux mentions obligatoires.

CE QUE LE CAHIER DES CHARGES DEMANDE (§5, « Should ») : « statuts, PV
d'AG, contrat de travail, mise en demeure, bail commercial — pre-remplis
par questions-reponses, CONFORMES AUX MENTIONS OBLIGATOIRES, export Word
et PDF ».

---------------------------------------------------------------------
LE QUESTIONNAIRE EST LU DANS L'ARTICLE, PAS ECRIT ICI.
---------------------------------------------------------------------

L'article 13 de l'AUSCGIE enumere ce que des statuts doivent contenir :
« 1° la forme de la societe ; 2° sa denomination ; 3° la nature et le
domaine de son activite... ». Chacune de ces mentions devient UNE
QUESTION, et la reponse devient UNE CLAUSE.

C'est la meme fonction que l'analyse de conformite
(conformite.mentions_obligatoires) : le document produit est donc, par
construction, celui que l'analyse validerait. Deux listes tenues
separement auraient fini par diverger — et l'outil aurait genere des
documents que son propre controle aurait rejetes.

Quand une revision change la liste, le questionnaire change avec elle,
sans qu'on ait rien a mettre a jour.

---------------------------------------------------------------------
AUCUN TEXTE N'EST INVENTE PAR UN MODELE.
---------------------------------------------------------------------

Le document est assemble MECANIQUEMENT : une clause par mention, portant
l'intitule lu dans la loi et la reponse saisie par l'utilisateur. Faire
rediger les clauses par un modele de langage produirait des phrases
plausibles et invérifiables dans un acte destine a etre signe — c'est
exactement le risque que ce produit existe pour ecarter.

L'utilisateur obtient donc un SQUELETTE FIDELE, a completer et a faire
relire, et non un acte pret a signer. Le document le dit lui-meme, sur
sa premiere page.
"""

from __future__ import annotations

import datetime
import io

from app.services.conformite import mentions_obligatoires

# Un document type, et l'article qui porte ses mentions obligatoires.
# LA LISTE N'EST PAS ICI : elle est lue dans l'article a chaque appel.
MODELES: dict[str, dict] = {
    "statuts_sarl": {
        "libelle": "Statuts de société à responsabilité limitée",
        "titre_document": "STATUTS",
        "sigle": "AUSCGIE",
        "numero": "13",
        "preambule": (
            "Les soussignés, propriétaires des parts ci-après créées, "
            "ont établi ainsi qu'il suit les statuts de la société."
        ),
    },
    "statuts_sa": {
        "libelle": "Statuts de société anonyme",
        "titre_document": "STATUTS DE SOCIÉTÉ ANONYME",
        "sigle": "AUSCGIE",
        "numero": "397",
        "preambule": (
            "Les soussignés, propriétaires des actions ci-après créées, "
            "ont établi ainsi qu'il suit les statuts de la société."
        ),
    },
    "statuts_cooperative": {
        "libelle": "Statuts de société coopérative",
        "titre_document": "STATUTS DE SOCIÉTÉ COOPÉRATIVE",
        "sigle": "AUSCOOP",
        "numero": "18",
        "preambule": (
            "Les soussignés ont établi ainsi qu'il suit les statuts de "
            "la société coopérative."
        ),
    },
}

AVERTISSEMENT = (
    "Ce document est un SQUELETTE produit automatiquement à partir des "
    "mentions que la loi impose. Il n'est ni rédigé ni relu par un "
    "professionnel : il doit être complété, adapté à votre situation et "
    "soumis à un conseil avant toute signature. ChatDocs OHADA est une "
    "aide à la recherche documentaire et ne garantit aucune conformité "
    "du document produit."
)


class GenerationRefusee(RuntimeError):
    """Le document ne peut pas être produit, et on dit pourquoi."""


def questionnaire(contenu_article: str) -> list[dict]:
    """Transforme les mentions obligatoires en questions.

    Une mention par question, dans l'ordre de l'article : c'est cet
    ordre que le juriste relira, et le changer rendrait la comparaison
    avec le texte de loi pénible pour rien.
    """
    return [
        {
            "cle": f"mention_{point['repere'].rstrip('°')}",
            "repere": point["repere"],
            "question": _en_question(point["libelle"]),
            "libelle_legal": point["libelle"],
        }
        for point in mentions_obligatoires(contenu_article)
    ]


def _en_question(libelle: str) -> str:
    """Présente l'intitulé légal comme libellé de champ.

    ON NE REFORMULE PAS, ON MET UNE MAJUSCULE. Une première tentative
    retirait le déterminant initial pour faire une question — « la forme
    de la société » devenait « Forme de la société ». Le procédé casse
    dès que l'intitulé en contient un second : « la nature et le domaine
    de son activité » donnait « Nature et LE domaine ».

    Surtout, réécrire un intitulé change ce que la loi demande. Le
    libellé légal fait un très bon libellé de champ tel quel, et il a
    l'avantage décisif de se comparer mot pour mot au texte de
    l'article pendant la relecture.
    """
    propre = libelle.strip()
    return propre[0].upper() + propre[1:] if propre else propre


def assembler(
    modele: dict,
    champs: list[dict],
    reponses: dict[str, str],
    version_corpus: str,
) -> dict:
    """Construit le document : une clause par mention obligatoire.

    UNE MENTION SANS REPONSE N'EST PAS OMISE. Elle figure au document
    avec un marqueur visible « [À COMPLÉTER] » : une clause absente
    passerait inaperçue à la relecture, alors qu'un trou saute aux yeux.
    C'est le même principe que « le doute profite au à vérifier » de
    l'analyse de conformité.
    """
    clauses = []
    for rang, champ in enumerate(champs, 1):
        reponse = (reponses.get(champ["cle"]) or "").strip()
        clauses.append(
            {
                "numero": rang,
                "repere": champ["repere"],
                "intitule": champ["libelle_legal"],
                "contenu": reponse or "[À COMPLÉTER]",
                "complete": bool(reponse),
            }
        )

    return {
        "titre": modele["titre_document"],
        "preambule": modele["preambule"],
        "base_legale": f"{modele['sigle']} — article {modele['numero']}",
        "version_corpus": version_corpus,
        "clauses": clauses,
        "restant_a_completer": sum(1 for c in clauses if not c["complete"]),
        "avertissement": AVERTISSEMENT,
    }


# ---------------------------------------------------------------------
# Exports
#
# DEUX FORMATS, DEUX USAGES. Le .docx est celui qu'on RETOUCHE — et un
# squelette est fait pour être retouché. Le PDF est celui qu'on
# TRANSMET. Ne livrer que le PDF obligerait à retaper le document.
# ---------------------------------------------------------------------


def en_docx(document: dict) -> bytes:
    """Document Word, retouchable."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    fichier = Document()

    titre = fichier.add_heading(document["titre"], level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    reference = fichier.add_paragraph()
    reference.alignment = WD_ALIGN_PARAGRAPH.CENTER
    piece = reference.add_run(
        f"Mentions obligatoires : {document['base_legale']}\n"
        f"Corpus : {document['version_corpus']}"
    )
    piece.italic = True
    piece.font.size = Pt(9)

    fichier.add_paragraph(document["preambule"])

    for clause in document["clauses"]:
        fichier.add_heading(
            f"Article {clause['numero']} — {clause['intitule']}", level=2
        )
        paragraphe = fichier.add_paragraph(clause["contenu"])
        if not clause["complete"]:
            # Le marqueur doit se voir dans Word aussi : un trou en gris
            # clair se relit comme du texte.
            paragraphe.runs[0].bold = True

    fichier.add_page_break()
    entete = fichier.add_paragraph()
    entete.add_run("Avertissement").bold = True
    note = fichier.add_paragraph(document["avertissement"])
    note.runs[0].font.size = Pt(8)

    tampon = io.BytesIO()
    fichier.save(tampon)
    return tampon.getvalue()


def en_pdf(document: dict, genere_le: datetime.date | None = None) -> bytes:
    """Document PDF, transmissible."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    from app.services.export_pdf import BLEU_NUIT, GRIS, OR, _echapper

    base = getSampleStyleSheet()
    styles = {
        "titre": ParagraphStyle(
            "titre",
            parent=base["Title"],
            fontName="Times-Bold",
            fontSize=17,
            textColor=BLEU_NUIT,
        ),
        "meta": ParagraphStyle(
            "meta",
            parent=base["Normal"],
            fontSize=8,
            alignment=TA_CENTER,
            textColor=GRIS,
        ),
        "preambule": ParagraphStyle(
            "preambule",
            parent=base["Normal"],
            fontSize=10,
            leading=15,
            alignment=TA_JUSTIFY,
            spaceBefore=5 * mm,
            spaceAfter=5 * mm,
        ),
        "intitule": ParagraphStyle(
            "intitule",
            parent=base["Normal"],
            fontName="Times-Bold",
            fontSize=11,
            textColor=BLEU_NUIT,
            spaceBefore=5 * mm,
            spaceAfter=1.5 * mm,
        ),
        "clause": ParagraphStyle(
            "clause",
            parent=base["Normal"],
            fontSize=10,
            leading=15,
            alignment=TA_JUSTIFY,
        ),
        "trou": ParagraphStyle(
            "trou",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#A5851F"),
        ),
        "avertissement": ParagraphStyle(
            "avertissement",
            parent=base["Normal"],
            fontSize=8,
            leading=12,
            textColor=GRIS,
        ),
    }

    tampon = io.BytesIO()
    pdf = SimpleDocTemplate(
        tampon,
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        title=document["titre"],
    )

    date = genere_le or datetime.date.today()
    elements = [
        Paragraph(_echapper(document["titre"]), styles["titre"]),
        Paragraph(
            f"Mentions obligatoires : {_echapper(document['base_legale'])}",
            styles["meta"],
        ),
        Paragraph(
            f"Corpus : {_echapper(document['version_corpus'])} — "
            f"édité le {date:%d/%m/%Y}",
            styles["meta"],
        ),
        Spacer(1, 3 * mm),
        HRFlowable(width="100%", color=OR, thickness=1.2),
        Paragraph(_echapper(document["preambule"]), styles["preambule"]),
    ]

    for clause in document["clauses"]:
        elements.append(
            Paragraph(
                f"Article {clause['numero']} — {_echapper(clause['intitule'])}",
                styles["intitule"],
            )
        )
        elements.append(
            Paragraph(
                _echapper(clause["contenu"]),
                styles["clause"] if clause["complete"] else styles["trou"],
            )
        )

    elements += [
        PageBreak(),
        Paragraph("Avertissement", styles["intitule"]),
        Paragraph(_echapper(document["avertissement"]), styles["avertissement"]),
    ]

    pdf.build(elements)
    return tampon.getvalue()
